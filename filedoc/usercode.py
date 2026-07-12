#!/usr/bin/env python3
"""Corezoid GIT Call entrypoint (lang=python, path=filedoc) — FILE → STRUCTURE + SAMPLE TEXT.

Feeds the Corezoid skill **dto-mf-file-doc**: extracts a STRUCTURE-FIRST, size-bounded text
representation of a client office file (xlsx/xls, docx, pdf, csv/tsv, xml, json, txt/md). The
process then hands `doc_text` to the inline LLM extractor (§3) → dto-mf-writer + dto-mf-graph.
This node ONLY reads/derives text — deterministic, no LLM, no Simulator calls. Honest: if the
format is unsupported or unreadable, `twin_error` is set, text is empty, nothing is invented.

BIG-FILE STRATEGY (#217):
  Never dump a huge file whole into the LLM. Instead:
    1. Extract the STRUCTURE first (sheet names + headers + row counts / TOC + page count /
       xml top-tags / json top-keys).
    2. SAMPLE representative rows/pages (head + evenly-spaced probes + tail), not blind first-N.
    3. SIZE-GUARD: if the file is bigger than maxKb (or a table/pdf exceeds sampleRows/maxPages)
       → mode="structural" (structure header + sample) instead of full text. Otherwise "full".
    4. Final text is always clipped to maxChars → the LLM token budget can never blow up.

Contract (git_call calls handle(data)):
  IN:  file_b64   (str|list) base64 of the file (preferred for small attachments), OR
       source_url (str)      http(s) URL the runner downloads (preferred for BIG files), OR
       path       (str)      local path the runner can read (dev/test)
       file_name  (str)      original name WITH extension (drives format detection)
       maxKb        (int)  size-guard threshold in KB           (from @mf-file-max-kb)
       maxChars     (int)  hard cap of LLM text                 (from @mf-file-doc-max-chars)
       sampleRows   (int)  rows to sample from big tables       (from @mf-file-sample-rows)
       maxPages     (int)  PDF pages actually read              (from @mf-file-max-pages)
       pdfKbPerPage (int)  PDF page weight for size-guard       (from @mf-est-pdf-kb-per-page)
  OUT: data["filedoc"] = {
         found, format, mode(full|structural), file_name, bytes, kb,
         text, n_chars, n_rows, n_pages, n_sheets,
         structure {..}, truncated(bool), sampled(bool), twin_error, lib_status
       }
  Mirrors text/file_name/kb/format/mode to top-level.
Окно git_call ≤30с / reply ≤1.4МБ: structural mode keeps both bounded regardless of file size.
"""
import os, sys, json, base64, tempfile, time, traceback, hashlib, re
from decimal import Decimal, InvalidOperation
from datetime import date, datetime
from urllib.request import urlretrieve

LIB = {}
try:
    import openpyxl
    LIB["openpyxl"] = getattr(openpyxl, "__version__", "?")
except Exception as e:
    openpyxl = None; LIB["openpyxl"] = "ERR:%s" % e
try:
    import docx
    LIB["python-docx"] = "ok"
except Exception as e:
    docx = None; LIB["python-docx"] = "ERR:%s" % e
try:
    import pypdf
    LIB["pypdf"] = getattr(pypdf, "__version__", "?")
except Exception as e:
    pypdf = None; LIB["pypdf"] = "ERR:%s" % e
try:
    import xml.etree.ElementTree as ET
    LIB["xml"] = "stdlib"
except Exception as e:
    ET = None; LIB["xml"] = "ERR:%s" % e

# Defaults (overridden by task-data resolved from env in the process).
DEF_MAX_CHARS = 16000
DEF_MAX_KB = 5120
DEF_SAMPLE_ROWS = 60
DEF_MAX_PAGES = 30
DEF_PDF_KB_PER_PAGE = 80


def _int(v, d):
    try:
        n = int(v)
        return n if n > 0 else d
    except Exception:
        return d


def _ext(name):
    return (os.path.splitext(name or "")[1] or "").lower().lstrip(".")


def _detect(path, name):
    ext = _ext(name) or _ext(path)
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception:
        head = b""
    if head[:4] == b"PK\x03\x04":  # zip container
        if ext == "docx":
            return "docx"
        if ext in ("xlsx", "xlsm"):
            return "xlsx"
        return ext or "zip"
    if head[:5] == b"%PDF-":
        return "pdf"
    hs = head.lstrip()
    if hs[:1] in (b"{", b"["):
        return "json"
    if hs[:5].lower() == b"<?xml" or hs[:1] == b"<":
        if ext not in ("html", "htm"):
            return "xml"
    if ext in ("xlsx", "xlsm", "xls", "docx", "doc", "pdf", "csv", "tsv",
               "txt", "md", "xml", "json"):
        return ext
    return ext or "unknown"


def _sample_indices(total, k):
    """Head + evenly-spaced probes + tail: representative, not blind first-N."""
    if total <= k:
        return list(range(total))
    head = max(1, k // 3)
    tail = max(1, k // 3)
    mid = k - head - tail
    idx = list(range(head))
    if mid > 0:
        span = total - head - tail
        step = span / float(mid + 1)
        for i in range(1, mid + 1):
            idx.append(head + int(round(step * i)))
    idx += list(range(total - tail, total))
    seen, out = set(), []
    for i in idx:
        if 0 <= i < total and i not in seen:
            seen.add(i); out.append(i)
    return out


def _materialize(data):
    """Return (local_path, source_name, tmp_created). BIG files come via source_url."""
    name = data.get("file_name") or data.get("fileName") or ""
    b64 = (data.get("file_b64") or data.get("content_b64") or data.get("base64")
           or data.get("file_b64_parts"))
    if b64:
        s = b64
        if isinstance(s, (list, tuple)):
            s = "".join(str(x) for x in s)
        if isinstance(s, (bytes, bytearray)):
            s = s.decode("ascii", "ignore")
        if s.startswith("data:") and "," in s:
            s = s.split(",", 1)[1]
        s = "".join(s.split())
        s += "=" * (-len(s) % 4)
        raw = base64.b64decode(s)
        fd, tmp = tempfile.mkstemp(suffix="_" + (name or "file"))
        with os.fdopen(fd, "wb") as f:
            f.write(raw)
        return tmp, (name or os.path.basename(tmp)), True
    url = (data.get("source_url") or data.get("url") or data.get("file_url") or "").strip()
    if url:
        base = name or os.path.basename(url.split("?", 1)[0]) or "file"
        fd, tmp = tempfile.mkstemp(suffix="_" + base)
        os.close(fd)
        urlretrieve(url, tmp)
        return tmp, base, True
    path = (data.get("path") or "").strip()
    if path:
        return path, (name or os.path.basename(path)), False
    raise RuntimeError("file_b64 or source_url required")


# ---------- per-format extractors: return (text, n_rows, n_pages, n_sheets, structure, sampled) ----------

def _xlsx(path, cfg):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets, lines, total_rows, sampled = [], [], 0, False
    for ws in wb.worksheets:
        mr = ws.max_row or 0
        header = None
        rows = []  # (idx, cells)
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [("" if c is None else str(c)) for c in row]
            if not any(x.strip() for x in cells):
                continue
            if header is None:
                header = cells
            rows.append((i, cells))
            if len(rows) > cfg["sampleRows"] * 4 and mr > cfg["sampleRows"]:
                # stop hoarding; we only need enough to sample from
                break
        nrows = mr if mr else len(rows)
        total_rows += nrows
        sheets.append({"name": ws.title, "rows": nrows,
                       "cols": len(header) if header else 0,
                       "headers": (header or [])[:20]})
        lines.append("# Лист: %s (рядків ~%d)" % (ws.title, nrows))
        if header:
            lines.append("Колонки: " + " | ".join(header[:20]))
        take = _sample_indices(len(rows), cfg["sampleRows"])
        if len(take) < len(rows):
            sampled = True
        for j in take:
            lines.append("\t".join(rows[j][1]).rstrip())
    wb.close()
    struct = {"sheets": sheets}
    return "\n".join(lines), total_rows, None, len(sheets), struct, sampled


def _csv(path, cfg):
    with open(path, "rb") as f:
        raw = f.read()
    txt = raw.decode("utf-8", "ignore")
    lines_all = txt.splitlines()
    # detect delimiter from first non-empty line
    delim = ","
    for ln in lines_all[:5]:
        if ln.strip():
            for d in ("\t", ";", ",", "|"):
                if d in ln:
                    delim = d; break
            break
    rows = [r for r in lines_all if r.strip()]
    header = rows[0] if rows else ""
    body = rows[1:]
    n = len(body)
    take = _sample_indices(n, cfg["sampleRows"])
    sampled = len(take) < n
    out = ["# Таблиця CSV (рядків %d, розділювач '%s')" % (n, "\\t" if delim == "\t" else delim)]
    if header:
        out.append("Заголовок: " + header)
    for j in take:
        out.append(body[j])
    struct = {"delimiter": ("\\t" if delim == "\t" else delim),
              "columns": header.split(delim)[:20], "rows": n}
    return "\n".join(out), n, None, None, struct, sampled


def _docx(path, cfg):
    d = docx.Document(path)
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    tbl_rows = []
    for tbl in d.tables:
        for r in tbl.rows:
            cells = [c.text.strip() for c in r.cells]
            if any(cells):
                tbl_rows.append("\t".join(cells))
    lines = ["# Документ DOCX (параграфів %d, таблиць %d, рядків таблиць %d)"
             % (len(paras), len(d.tables), len(tbl_rows))]
    lines += paras
    if tbl_rows:
        take = _sample_indices(len(tbl_rows), cfg["sampleRows"])
        lines.append("# Таблиці (семпл):")
        for j in take:
            lines.append(tbl_rows[j])
    sampled = len(tbl_rows) > cfg["sampleRows"]
    struct = {"paragraphs": len(paras), "tables": len(d.tables), "table_rows": len(tbl_rows)}
    return "\n".join(lines), (len(tbl_rows) or None), None, None, struct, sampled


def _pdf(path, cfg):
    r = pypdf.PdfReader(path)
    n_pages = len(r.pages)
    take = _sample_indices(n_pages, cfg["maxPages"])
    sampled = len(take) < n_pages
    out = ["# PDF (сторінок %d, прочитано %d)" % (n_pages, len(take))]
    for pi in take:
        try:
            t = (r.pages[pi].extract_text() or "").strip()
        except Exception:
            t = ""
        if t:
            out.append("## Сторінка %d\n%s" % (pi + 1, t))
        if sum(len(x) for x in out) > cfg["maxChars"] * 2:
            break
    struct = {"pages": n_pages, "pages_read": len(take)}
    return "\n".join(out), None, n_pages, None, struct, sampled


def _xml(path, cfg):
    # Streaming: count elements, capture top-level tags; sample serialized head.
    counts, order, total = {}, [], 0
    root_tag = None
    try:
        for ev, el in ET.iterparse(path, events=("start",)):
            total += 1
            tag = el.tag.split("}")[-1]
            if root_tag is None:
                root_tag = tag
            if tag not in counts:
                counts[tag] = 0; order.append(tag)
            counts[tag] += 1
            if total > 50000:
                break
    except Exception:
        pass
    with open(path, "rb") as f:
        head = f.read(cfg["maxChars"]).decode("utf-8", "ignore")
    top = [(t, counts[t]) for t in order[:30]]
    out = ["# XML (root=%s, вузлів ~%d)" % (root_tag or "?", total)]
    out.append("Теги: " + ", ".join("%s×%d" % (t, c) for t, c in top))
    out.append("# Фрагмент:")
    out.append(head)
    struct = {"root": root_tag, "nodes": total, "tags": dict(top)}
    return "\n".join(out), total, None, None, struct, (total > 30)


def _json(path, cfg):
    with open(path, "rb") as f:
        raw = f.read().decode("utf-8", "ignore")
    try:
        obj = json.loads(raw)
    except Exception:
        # too big / invalid to fully parse — fall back to head text
        return ("# JSON (не розпарсено повністю)\n" + raw[:cfg["maxChars"]],
                None, None, None, {"parsed": False}, True)
    out, struct, sampled = [], {}, False
    if isinstance(obj, list):
        n = len(obj)
        struct = {"type": "array", "length": n}
        out.append("# JSON масив (елементів %d)" % n)
        take = _sample_indices(n, max(5, cfg["sampleRows"] // 6))
        sampled = len(take) < n
        if obj and isinstance(obj[0], dict):
            struct["item_keys"] = list(obj[0].keys())[:30]
            out.append("Ключі елемента: " + ", ".join(struct["item_keys"]))
        for j in take:
            out.append(json.dumps(obj[j], ensure_ascii=False)[:800])
    elif isinstance(obj, dict):
        keys = list(obj.keys())
        struct = {"type": "object", "keys": keys[:40]}
        out.append("# JSON обʼєкт (ключів %d)" % len(keys))
        out.append("Ключі: " + ", ".join(keys[:40]))
        out.append(json.dumps(obj, ensure_ascii=False)[:cfg["maxChars"]])
    else:
        out.append(str(obj)[:cfg["maxChars"]])
    return "\n".join(out), None, None, None, struct, sampled


def _txt(path, cfg):
    with open(path, "rb") as f:
        raw = f.read()
    txt = raw.decode("utf-8", "ignore")
    mc = cfg["maxChars"]
    if len(txt) <= mc:
        return txt, None, None, None, {"chars": len(txt)}, False
    # head + tail (both ends), not blind first-N
    head = txt[: int(mc * 0.6)]
    tail = txt[-int(mc * 0.35):]
    out = head + "\n\n…[СЕРЕДИНУ ПРОПУЩЕНО, файл великий]…\n\n" + tail
    return out, None, None, None, {"chars": len(txt)}, True


# =========================================================================================
# mode="load" — ДЕТЕРМІНОВАНИЙ ПОРЯДКОВИЙ ПАРСЕР (степпер) по MappingContract → entities[].
# Жодного AI. Числа/гроші — ТІЛЬКИ Decimal/строка (НЕ float). Резюмівний cursor {sheet,offset,tableIdx}.
# Контракт adapter (OUT.entities[]):
#   {class, title, code?, data{}, accounts:[{name, amount(str Decimal), currency, kind, unit}],
#    links:[{fromValue, toFileTable, toKeyCol, edge_type}]}
# MappingContract (IN):
#   {locale: {decimal: "," | "."}, ref_template (глобальний, опц.),
#    tables: [{sheet? | tableIdx?, name?, header_row(деф 1),
#      rowEntity: {class, ref_template, title_col, skip_if?},
#      columns: [{col, role: key|attr|account|link|ignore,
#                 account_name?, unit?, kind?, currency?,          # role=account
#                 to_table?, to_key?, edge_type?,                  # role=link
#                 date? (bool)}]}]}
# Вікно: chunk_rows (деф 200) АБО ~25с — що раніше. cursor.offset += видано в поточній таблиці.
# =========================================================================================

LOAD_TIME_BUDGET_S = 25.0
DEF_CHUNK_ROWS = 200
_NUM_CLEAN_RE = re.compile(r"[^\d,.\-]")


def _to_json(v):
    """Приймає dict АБО JSON-строку (Corezoid api_rpc часто передає JSON-строкою)."""
    if isinstance(v, (dict, list)):
        return v
    if isinstance(v, str) and v.strip():
        try:
            return json.loads(v)
        except ValueError:
            return None
    return None


def _norm_number(raw, decimal_sep="."):
    """Строка/число → канонічна Decimal-строка. НІКОЛИ float. Локаль з contract.
    Повертає (decimal_str|None, ok:bool). Порожнє → ('0', True)? Ні — порожнє → (None,False)
    щоб не вигадувати; викликаючий вирішує (account з None-сумою → пропускаємо суму)."""
    if raw is None:
        return None, False
    if isinstance(raw, bool):
        return None, False
    if isinstance(raw, int):
        return str(raw), True
    if isinstance(raw, Decimal):
        return format(raw, "f"), True
    if isinstance(raw, float):
        # float на вході (напр. openpyxl без int) — через str, не через binary float
        try:
            return format(Decimal(repr(raw)), "f"), True
        except (InvalidOperation, ValueError):
            return None, False
    s = str(raw).strip()
    if not s:
        return None, False
    # прибрати валютні символи/пробіли-роздільники тисяч, лишити цифри/,/./-
    s = _NUM_CLEAN_RE.sub("", s)
    if not s or s in ("-", ".", ","):
        return None, False
    if decimal_sep == ",":
        # кома — десятковий; крапка — роздільник тисяч
        s = s.replace(".", "").replace(",", ".")
    else:
        # крапка — десятковий; кома — роздільник тисяч
        s = s.replace(",", "")
    try:
        return format(Decimal(s), "f"), True
    except (InvalidOperation, ValueError):
        return None, False


def _norm_date(raw):
    """Дата → isoformat (YYYY-MM-DD). Приймає datetime/date/строку. None якщо не дата."""
    if raw is None or raw == "":
        return None
    if isinstance(raw, datetime):
        return raw.date().isoformat()
    if isinstance(raw, date):
        return raw.isoformat()
    s = str(raw).strip()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y", "%Y/%m/%d", "%m/%d/%Y"):
        try:
            return datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue
    return s  # лишаємо як є (честно), не вигадуємо


def _slug(v):
    """Детермінований slug для ref (стабільний → ідемпотентність)."""
    s = str(v or "").strip().lower()
    s = re.sub(r"[^\w]+", "-", s, flags=re.UNICODE).strip("-")
    return s or hashlib.sha1(str(v).encode("utf-8")).hexdigest()[:10]


def _make_ref(tmpl, key_value, cls):
    """ref_template з плейсхолдерами {key} / {class} / {slug}. Без шаблону → class-slug."""
    kv = _slug(key_value)
    if tmpl:
        return (tmpl.replace("{key}", str(key_value or ""))
                    .replace("{slug}", kv)
                    .replace("{class}", _slug(cls)))
    return "%s-%s" % (_slug(cls), kv)


def _xlsx_rows_window(path, sheet_name, header_row, offset, chunk_rows, t0):
    """Стрім рядків одного листа: (headers, [(row_cells)], produced, exhausted).
    read_only + iter_rows(values_only). Вікно: chunk_rows АБО час."""
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = None
    if sheet_name:
        for w in wb.worksheets:
            if w.title == sheet_name:
                ws = w; break
        if ws is None:
            wb.close()
            raise KeyError("sheet %r not found" % sheet_name)
    else:
        ws = wb.worksheets[0]
    headers = None
    out_rows = []
    produced = 0
    exhausted = True
    # min_row у openpyxl 1-based; header_row-й рядок = заголовок, дані з header_row+1
    start_data = header_row + 1 + offset
    ri = 0
    for row in ws.iter_rows(min_row=1, values_only=True):
        ri += 1
        if ri == header_row:
            headers = [("" if c is None else str(c).strip()) for c in row]
            continue
        if ri < start_data:
            continue
        cells = list(row)
        # порожній рядок — пропускаємо (не рахуємо в offset споживання даних? рахуємо як прочитаний)
        if not any(("" if c is None else str(c)).strip() for c in cells):
            continue
        out_rows.append(cells)
        produced += 1
        if produced >= chunk_rows or (time.time() - t0) >= LOAD_TIME_BUDGET_S:
            exhausted = False
            break
    else:
        exhausted = True
    if headers is None:
        headers = []
    wb.close()
    return headers, out_rows, produced, exhausted


def _docx_rows_window(path, table_idx, header_row, offset, chunk_rows, t0):
    """Рядки таблиці docx (одномоментно читаємо всю таблицю — вони малі; вікно на видачу)."""
    d = docx.Document(path)
    if table_idx >= len(d.tables):
        raise KeyError("tableIdx %d out of range (%d tables)" % (table_idx, len(d.tables)))
    tbl = d.tables[table_idx]
    all_rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
    if not all_rows:
        return [], [], 0, True
    headers = all_rows[header_row - 1] if len(all_rows) >= header_row else all_rows[0]
    body = all_rows[header_row:]  # після заголовка
    window = body[offset: offset + chunk_rows]
    exhausted = (offset + len(window)) >= len(body)
    return headers, window, len(window), exhausted


def _apply_contract(headers, row_cells, tbl_contract, locale_sep, acc, sums, unmatched):
    """Один рядок → entity по MappingContract (детерміновано). Оновлює acc/sums/unmatched.
    Повертає (entity|None, twin_error|None)."""
    # індекс колонки по імені (як у файлі)
    col_idx = {h: i for i, h in enumerate(headers)}

    def cell(colname):
        i = col_idx.get(colname)
        if i is None or i >= len(row_cells):
            return None
        v = row_cells[i]
        return v

    re_ = tbl_contract["rowEntity"]
    cls = re_["class"]
    ref_tmpl = re_.get("ref_template") or tbl_contract.get("_ref_template_global")
    title_col = re_.get("title_col")

    # skip_if: {col, equals} — детерміноване пропускання
    skip = re_.get("skip_if")
    if skip:
        sv = cell(skip.get("col"))
        if str(sv).strip() == str(skip.get("equals", "")).strip():
            return None, None

    entity = {"class": cls, "title": "", "data": {}, "accounts": [], "links": []}
    key_value = None

    for c in tbl_contract["columns"]:
        colname = c["col"]
        if colname not in col_idx:
            # колонка контракту відсутня у файлі — чесна помилка
            return None, "колонка '%s' відсутня у файлі (headers=%s)" % (colname, headers)
        role = c.get("role", "ignore")
        raw = cell(colname)

        if role == "ignore":
            continue
        if role == "key":
            key_value = ("" if raw is None else str(raw).strip())
            entity["code"] = key_value
            entity["data"][colname] = key_value
        elif role == "attr":
            if c.get("date"):
                entity["data"][colname] = _norm_date(raw)
            else:
                entity["data"][colname] = ("" if raw is None else str(raw).strip())
        elif role == "account":
            amt, ok = _norm_number(raw, locale_sep)
            if ok:
                a = {"name": c.get("account_name") or colname,
                     "amount": amt,   # СТРОКА Decimal
                     "currency": c.get("currency") or "one",
                     "kind": c.get("kind") or "fact",
                     "unit": c.get("unit") or ""}
                entity["accounts"].append(a)
                # акумулюємо суму для baseline-звірки
                prev = sums.get(a["name"], Decimal("0"))
                sums[a["name"]] = prev + Decimal(amt)
        elif role == "link":
            fv = ("" if raw is None else str(raw).strip())
            if fv:
                entity["links"].append({
                    "fromValue": fv,
                    "toFileTable": c.get("to_table"),
                    "toKeyCol": c.get("to_key"),
                    "edge_type": c.get("edge_type") or "related",
                })

    if title_col is not None:
        tv = cell(title_col)
        entity["title"] = ("" if tv is None else str(tv).strip())
    if not entity["title"]:
        entity["title"] = entity.get("code") or ""

    if not key_value:
        unmatched[0] += 1  # рядок без ключа — рахуємо, не мовчимо
        # все одно віддаємо (adapter може дедупити), але позначимо
        entity["_no_key"] = True

    ref = _make_ref(ref_tmpl, key_value or entity["title"], cls)
    entity["ref"] = ref
    acc["count"] += 1
    return entity, None


def _load(data):
    """mode=load: віддає наступний чанк entities[] + cursor. Резюмівно, детерміновано."""
    contract = _to_json(data.get("mapping_contract"))
    if not isinstance(contract, dict):
        return {"twin_error": "mapping_contract required (dict|json)", "entities": [],
                "cursor": {}, "done": True, "count": 0}
    cursor = _to_json(data.get("cursor")) or {}
    if not isinstance(cursor, dict):
        cursor = {}
    chunk_rows = _int(data.get("chunk_rows"), DEF_CHUNK_ROWS)
    locale = contract.get("locale") or {}
    locale_sep = locale.get("decimal", ".")
    ref_tmpl_global = contract.get("ref_template")
    tables = contract.get("tables") or []
    if not tables:
        return {"twin_error": "mapping_contract.tables empty", "entities": [],
                "cursor": {}, "done": True, "count": 0}

    # прокинути глобальний ref_template у кожну таблицю (fallback)
    for t in tables:
        t.setdefault("_ref_template_global", ref_tmpl_global)

    path, name, tmp_created = None, None, False
    t0 = time.time()
    try:
        path, name, tmp_created = _materialize(data)
        fmt = _detect(path, name)

        tbl_i = int(cursor.get("tableIdx", 0))
        offset = int(cursor.get("offset", 0))

        acc = {"count": 0}
        sums = {}          # accountName -> Decimal (у ЦЬОМУ чанку)
        source_counts = {} # sheet/table -> produced rows (у ЦЬОМУ чанку)
        unmatched = [0]
        entities = []
        twin_error = None

        # пройти таблиці контракту, поки не наберемо вікно або не вичерпаємо всі
        done = False
        while tbl_i < len(tables):
            tc = tables[tbl_i]
            header_row = _int(tc.get("header_row"), 1)
            if fmt in ("xlsx", "xlsm"):
                if openpyxl is None:
                    raise RuntimeError("openpyxl unavailable")
                sheet_name = tc.get("sheet")
                headers, rows, produced, exhausted = _xlsx_rows_window(
                    path, sheet_name, header_row, offset, chunk_rows - acc["count"], t0)
                sc_key = sheet_name or "sheet0"
            elif fmt == "docx":
                if docx is None:
                    raise RuntimeError("python-docx unavailable")
                table_idx = _int(tc.get("tableIdx"), 0) if tc.get("tableIdx") is not None else 0
                headers, rows, produced, exhausted = _docx_rows_window(
                    path, table_idx, header_row, offset, chunk_rows - acc["count"], t0)
                sc_key = "table%d" % table_idx
            else:
                raise RuntimeError("mode=load: формат '%s' не підтримується (xlsx/docx)" % fmt)

            for rc in rows:
                ent, err = _apply_contract(headers, rc, tc, locale_sep, acc, sums, unmatched)
                if err:
                    twin_error = err
                    break
                if ent is not None:
                    entities.append(ent)
            source_counts[sc_key] = source_counts.get(sc_key, 0) + produced

            if twin_error:
                break

            if exhausted:
                # ця таблиця вичерпана → наступна з offset 0
                tbl_i += 1
                offset = 0
                if tbl_i >= len(tables):
                    done = True
                    break
                # якщо вже набрали вікно — стоп (наступна таблиця у наступному виклику)
                if acc["count"] >= chunk_rows or (time.time() - t0) >= LOAD_TIME_BUDGET_S:
                    break
            else:
                # таблиця не вичерпана → зсуваємо offset у ній, вікно закрите
                offset += produced
                break

        next_cursor = {"tableIdx": tbl_i, "offset": offset}
        out = {
            "entities": entities,
            "cursor": (next_cursor if not done else {"tableIdx": tbl_i, "offset": 0, "done": True}),
            "done": bool(done),
            "format": fmt,
            "count": len(entities),
            "source_counts": source_counts,
            "sums": {k: format(v, "f") for k, v in sums.items()},  # СТРОКИ Decimal
            "unmatched": unmatched[0],
            "twin_error": twin_error,
        }
        return out
    except Exception as e:
        return {"twin_error": "%s: %s" % (type(e).__name__, str(e)[:300]),
                "twin_trace": traceback.format_exc()[-1000:],
                "entities": [], "cursor": {}, "done": True, "count": 0}
    finally:
        if tmp_created and path:
            try:
                os.unlink(path)
            except Exception:
                pass


def handle(data, context=None):
    # ---- mode=load: окрема детермінована гілка (степпер по MappingContract) ----
    if (data.get("mode") or "").lower() == "load":
        res = _load(data)
        for k in ("file_b64", "content_b64", "base64", "file_b64_parts"):
            if k in data:
                try:
                    del data[k]
                except Exception:
                    data[k] = ""
        data["entities"] = res.get("entities", [])
        data["cursor"] = res.get("cursor", {})
        data["done"] = bool(res.get("done", True))
        data["format"] = res.get("format", "")
        data["count"] = res.get("count", 0)
        data["source_counts"] = res.get("source_counts", {})
        data["sums"] = res.get("sums", {})
        data["unmatched"] = res.get("unmatched", 0)
        data["twin_error"] = res.get("twin_error")
        if res.get("twin_trace"):
            data["twin_trace"] = res["twin_trace"]
        return data

    cfg = {
        "maxChars": _int(data.get("maxChars"), DEF_MAX_CHARS),
        "maxKb": _int(data.get("maxKb"), DEF_MAX_KB),
        "sampleRows": _int(data.get("sampleRows"), DEF_SAMPLE_ROWS),
        "maxPages": _int(data.get("maxPages"), DEF_MAX_PAGES),
        "pdfKbPerPage": _int(data.get("pdfKbPerPage"), DEF_PDF_KB_PER_PAGE),
    }
    out = {"found": False, "format": None, "mode": "full", "file_name": None,
           "bytes": 0, "kb": 0, "text": "", "n_chars": 0, "n_rows": None,
           "n_pages": None, "n_sheets": None, "structure": {}, "truncated": False,
           "sampled": False, "twin_error": None, "lib_status": LIB}
    t0 = time.time()
    tmp_created = False
    path = None
    try:
        path, name, tmp_created = _materialize(data)
        out["file_name"] = name
        try:
            out["bytes"] = os.path.getsize(path)
            out["kb"] = int(round(out["bytes"] / 1024.0))
        except Exception:
            pass
        fmt = _detect(path, name)
        out["format"] = fmt

        text, nrows, npages, nsheets, struct, sampled = "", None, None, None, {}, False
        if fmt in ("xlsx", "xlsm"):
            if openpyxl is None:
                raise RuntimeError("openpyxl unavailable")
            text, nrows, npages, nsheets, struct, sampled = _xlsx(path, cfg)
        elif fmt in ("csv", "tsv"):
            text, nrows, npages, nsheets, struct, sampled = _csv(path, cfg)
        elif fmt == "docx":
            if docx is None:
                raise RuntimeError("python-docx unavailable")
            text, nrows, npages, nsheets, struct, sampled = _docx(path, cfg)
        elif fmt == "pdf":
            if pypdf is None:
                raise RuntimeError("pypdf unavailable")
            text, nrows, npages, nsheets, struct, sampled = _pdf(path, cfg)
        elif fmt == "xml":
            if ET is None:
                raise RuntimeError("xml unavailable")
            text, nrows, npages, nsheets, struct, sampled = _xml(path, cfg)
        elif fmt == "json":
            text, nrows, npages, nsheets, struct, sampled = _json(path, cfg)
        elif fmt in ("txt", "md"):
            text, nrows, npages, nsheets, struct, sampled = _txt(path, cfg)
        else:
            out["twin_error"] = ("формат '%s' не підтримується (xlsx/xls/docx/pdf/csv/tsv/"
                                 "xml/json/txt/md); текст не витягнуто" % fmt)

        # ----- SIZE-GUARD: decide full vs structural -----
        est_kb = out["kb"]
        if fmt == "pdf" and npages:
            est_kb = max(est_kb, npages * cfg["pdfKbPerPage"])
        over_size = est_kb > cfg["maxKb"]
        text = (text or "").strip()
        truncated = False
        if len(text) > cfg["maxChars"]:
            text = text[: cfg["maxChars"]]
            truncated = True
        if over_size or sampled or truncated:
            out["mode"] = "structural"
        out["truncated"] = truncated
        out["sampled"] = bool(sampled)
        out["text"] = text
        out["n_chars"] = len(text)
        out["n_rows"] = nrows
        out["n_pages"] = npages
        out["n_sheets"] = nsheets
        out["structure"] = struct
        out["found"] = bool(text)
        if not text and not out["twin_error"]:
            out["twin_error"] = "текст не витягнуто (порожній або нечитабельний файл)"
    except Exception as e:
        out["twin_error"] = "%s: %s" % (type(e).__name__, e)
        out["trace"] = traceback.format_exc()[-800:]
    finally:
        if tmp_created and path:
            try:
                os.unlink(path)
            except Exception:
                pass
    out["ms"] = int((time.time() - t0) * 1000)
    for k in ("file_b64", "content_b64", "base64", "file_b64_parts"):
        if k in data:
            try:
                del data[k]
            except Exception:
                data[k] = ""
    data["filedoc"] = out
    data["doc_text"] = out["text"]
    data["file_name_out"] = out["file_name"]
    data["found"] = out["found"]
    data["twin_error"] = out["twin_error"]
    data["fd_format"] = out["format"]
    data["fd_mode"] = out["mode"]
    data["fd_kb"] = out["kb"]
    data["fd_nchars"] = out["n_chars"]
    data["fd_truncated"] = "1" if out["truncated"] else "0"
    data["fd_lib"] = out["lib_status"]
    return data


def usercode(data, context=None):
    return handle(data, context)


if __name__ == "__main__":
    import glob
    args = sys.argv[1:]
    if not args:
        args = sorted(glob.glob(os.path.join(os.path.dirname(__file__), "..", "..", "..",
                      "sprint-mf-2026-07-06", "fixtures", "docs", "*")))
    for a in args:
        r = handle({"path": a, "file_name": os.path.basename(a),
                    "maxChars": 4000, "maxKb": 200, "sampleRows": 20, "maxPages": 5})
        d = r["filedoc"]
        print("== %s | fmt=%s mode=%s kb=%s chars=%s rows=%s pages=%s sheets=%s "
              "trunc=%s sampled=%s err=%s" % (
                  d["file_name"], d["format"], d["mode"], d["kb"], d["n_chars"],
                  d["n_rows"], d["n_pages"], d["n_sheets"], d["truncated"],
                  d["sampled"], d["twin_error"]))
        print("struct:", json.dumps(d["structure"], ensure_ascii=False)[:300])
        print(d["text"][:400])
        print("----")
